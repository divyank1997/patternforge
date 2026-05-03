from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────────────
VIOLET  = RGBColor(0x7C, 0x3A, 0xED)
DARK    = RGBColor(0x1E, 0x1E, 0x2E)
GRAY    = RGBColor(0x6B, 0x72, 0x80)

def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=18, bold=True, color=VIOLET)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '7C3AED')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=DARK)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=GRAY)
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '7C3AED')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=10)
            if r_idx % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F3F4F6')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    # col widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PATTERNFORGE")
set_font(run, size=36, bold=True, color=VIOLET)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Business Requirements Document")
set_font(run, size=20, bold=True, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered 3D Pattern Generation Platform")
set_font(run, size=14, color=GRAY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Version 1.0   •   {datetime.date.today().strftime('%B %d, %Y')}")
set_font(run, size=11, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Author: Divyank Srivastava   •   Confidential")
set_font(run, size=11, color=GRAY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. DOCUMENT CONTROL
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Document Control")

heading2("1.1 Revision History")
add_table(
    ["Version", "Date", "Author", "Description"],
    [
        ["1.0", datetime.date.today().strftime("%Y-%m-%d"), "Divyank Srivastava", "Initial BRD — full platform scope"],
    ],
    [0.8, 1.2, 1.8, 2.7]
)

heading2("1.2 Document Purpose")
body("This Business Requirements Document (BRD) defines the functional and non-functional requirements for PatternForge — an AI-powered 3D generative pattern platform. It serves as the single source of truth for product, engineering, and design decisions throughout the project lifecycle.")

heading2("1.3 Intended Audience")
for a in ["Engineering Team", "Product Manager", "UX/UI Designers", "QA Engineers", "Potential Investors / Recruiters"]:
    bullet(a)

# ══════════════════════════════════════════════════════════════════════════════
# 2. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("2. Executive Summary")

heading2("2.1 Project Overview")
body("PatternForge is a cloud-native SaaS platform that allows users to generate, render, search, and share 3D generative patterns using natural language prompts powered by Claude AI. The platform targets creative professionals, designers, developers, and researchers who need rapid, customisable, high-quality visual patterns for use in UI design, textile, architecture, game development, and digital art.")

heading2("2.2 Problem Statement")
body("Creating complex 3D generative patterns currently requires:")
for p_text in [
    "Deep expertise in parametric design tools (Blender, Houdini, TouchDesigner)",
    "Significant time investment — hours or days to iterate on a single design",
    "No accessible platform for discovering, remixing, and sharing generative patterns",
    "No AI-native tool that translates human intent into structured 3D parameters",
]:
    bullet(p_text)

heading2("2.3 Proposed Solution")
body("PatternForge solves this by combining:")
for s in [
    "Claude AI — translates a text description into structured pattern parameters",
    "React Three Fiber — renders the pattern in real-time 3D in the browser",
    "OpenSearch — full-text and vector-similarity search across the pattern library",
    "Micro Frontend + Microservices architecture — independently scalable, modular, production-grade",
]:
    bullet(s)

heading2("2.4 Success Metrics")
add_table(
    ["Metric", "Target", "Timeframe"],
    [
        ["User Registrations", "500+", "Month 3"],
        ["Patterns Generated", "5,000+", "Month 3"],
        ["Average Pattern Generation Time", "< 3 seconds", "Launch"],
        ["System Uptime", "99.9%", "Ongoing"],
        ["API Response Time (p99)", "< 200ms", "Launch"],
        ["Search Result Relevance (NDCG)", "> 0.85", "Month 2"],
    ],
    [2.5, 1.5, 1.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. STAKEHOLDERS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("3. Stakeholders & User Personas")

heading2("3.1 Primary Stakeholders")
add_table(
    ["Stakeholder", "Role", "Interest"],
    [
        ["Divyank Srivastava", "Founder / Lead Engineer", "Technical excellence, portfolio showcase"],
        ["End Users", "Designers & Developers", "Fast, intuitive pattern generation"],
        ["FAANG Recruiters", "Technical Evaluators", "Architecture quality, code depth"],
    ],
    [2.0, 2.0, 2.5]
)

heading2("3.2 User Personas")

heading3("Persona 1 — The Creative Designer")
for item in [
    "Age: 25–35 | Role: UI/UX Designer at a tech startup",
    "Goal: Generate unique background patterns for app designs without coding",
    "Pain Point: Current tools require Blender skills or expensive licenses",
    "PatternForge Value: Describe a pattern in English → get a 3D-rendered result in seconds",
]:
    bullet(item)

heading3("Persona 2 — The Frontend Developer")
for item in [
    "Age: 22–32 | Role: Frontend Engineer",
    "Goal: Export pattern parameters to use in Three.js or CSS projects",
    "Pain Point: No structured format to define procedural patterns programmatically",
    "PatternForge Value: JSON parameter export, fork & remix public patterns",
]:
    bullet(item)

heading3("Persona 3 — The Researcher / Generative Artist")
for item in [
    "Age: 28–45 | Role: Academic / Digital Artist",
    "Goal: Explore and publish novel generative pattern techniques",
    "Pain Point: No community platform for sharing generative work with source parameters",
    "PatternForge Value: Public gallery, likes, forks, version history, searchable by style",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 4. SCOPE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("4. Project Scope")

heading2("4.1 In Scope")
for item in [
    "User authentication (register, login, JWT refresh, logout)",
    "AI-powered pattern generation from natural language prompts (Claude API)",
    "Real-time 3D pattern rendering in browser (React Three Fiber / Three.js)",
    "Pattern parameter editor with live preview",
    "Pattern save, fork, version history",
    "Public gallery with infinite scroll",
    "Full-text and vector similarity search (OpenSearch)",
    "User dashboard with analytics (patterns created, likes received)",
    "Micro Frontend architecture (shell + 4 MFEs)",
    "Microservices backend (6 services)",
    "Infrastructure as Code (AWS CDK)",
    "CI/CD pipeline (GitHub Actions)",
]:
    bullet(item)

heading2("4.2 Out of Scope (v1.0)")
for item in [
    "Mobile native apps (iOS / Android)",
    "Real-time collaboration / multiplayer editing",
    "Monetisation / billing / subscription tiers",
    "3D model export (OBJ, GLTF) — planned for v1.1",
    "Social features (comments, follows) — planned for v1.2",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 5. FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("5. Functional Requirements")

heading2("5.1 Authentication Module (mfe-auth + auth-service)")
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ["FR-AUTH-01", "User can register with email, username, display name, and password", "Must Have"],
        ["FR-AUTH-02", "Passwords stored as bcrypt hash (cost factor 12)", "Must Have"],
        ["FR-AUTH-03", "System issues JWT access token (15 min) + refresh token (7 days) on login", "Must Have"],
        ["FR-AUTH-04", "Refresh token rotation — old token invalidated on each refresh", "Must Have"],
        ["FR-AUTH-05", "GET /auth/me returns current user from valid access token", "Must Have"],
        ["FR-AUTH-06", "Logout invalidates the refresh token in the database", "Must Have"],
        ["FR-AUTH-07", "Duplicate email or username returns 409 Conflict with clear message", "Must Have"],
    ],
    [1.0, 4.5, 1.0]
)

heading2("5.2 AI Pattern Generation (mfe-editor + ai-service)")
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ["FR-AI-01", "User submits a text prompt describing a desired pattern", "Must Have"],
        ["FR-AI-02", "Claude API translates the prompt into structured PatternParameters JSON", "Must Have"],
        ["FR-AI-03", "Generated parameters include: style, complexity, scale, rotation, symmetry, color palette, seed", "Must Have"],
        ["FR-AI-04", "AI response streamed to client for perceived speed", "Should Have"],
        ["FR-AI-05", "User can regenerate with the same prompt (different seed)", "Must Have"],
        ["FR-AI-06", "AI explanation of the pattern is displayed alongside the 3D render", "Should Have"],
    ],
    [1.0, 4.5, 1.0]
)

heading2("5.3 3D Pattern Renderer (mfe-editor)")
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ["FR-3D-01", "Pattern rendered in real-time using React Three Fiber", "Must Have"],
        ["FR-3D-02", "User can orbit, zoom, and pan the 3D view", "Must Have"],
        ["FR-3D-03", "Parameter sliders update the render in real-time without page reload", "Must Have"],
        ["FR-3D-04", "Supported pattern styles: geometric, organic, fractal, noise, parametric", "Must Have"],
        ["FR-3D-05", "Color palette picker with primary, secondary, accent, and background colors", "Must Have"],
        ["FR-3D-06", "Screenshot / thumbnail auto-captured on save", "Should Have"],
    ],
    [1.0, 4.5, 1.0]
)

heading2("5.4 Pattern Management (pattern-service)")
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ["FR-PAT-01", "Authenticated user can save a pattern with title and description", "Must Have"],
        ["FR-PAT-02", "Pattern stored in DynamoDB with full parameter JSON", "Must Have"],
        ["FR-PAT-03", "Each save creates a new version; version history is browsable", "Should Have"],
        ["FR-PAT-04", "User can make a pattern public or private", "Must Have"],
        ["FR-PAT-05", "User can fork any public pattern to their own library", "Must Have"],
        ["FR-PAT-06", "Pattern metadata indexed in OpenSearch on save", "Must Have"],
    ],
    [1.0, 4.5, 1.0]
)

heading2("5.5 Search & Discovery (mfe-gallery + search-service)")
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ["FR-SRCH-01", "Full-text search across pattern title, description, tags, and prompt", "Must Have"],
        ["FR-SRCH-02", "Filter by style, color palette, and tags", "Must Have"],
        ["FR-SRCH-03", "Vector similarity search — find visually similar patterns", "Should Have"],
        ["FR-SRCH-04", "Gallery displays patterns in responsive masonry grid", "Must Have"],
        ["FR-SRCH-05", "Infinite scroll pagination using cursor-based pagination", "Must Have"],
        ["FR-SRCH-06", "Sort by: newest, most liked, trending", "Should Have"],
    ],
    [1.0, 4.5, 1.0]
)

heading2("5.6 Dashboard (mfe-dashboard)")
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ["FR-DASH-01", "User sees total patterns created, total likes, and public/private split", "Must Have"],
        ["FR-DASH-02", "Activity chart showing patterns created over time (Recharts)", "Should Have"],
        ["FR-DASH-03", "Quick access to recent patterns with thumbnail preview", "Must Have"],
        ["FR-DASH-04", "Profile edit (display name, avatar URL)", "Should Have"],
    ],
    [1.0, 4.5, 1.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("6. Non-Functional Requirements")

add_table(
    ["ID", "Category", "Requirement", "Target"],
    [
        ["NFR-01", "Performance", "API Gateway response time (p95)", "< 100ms"],
        ["NFR-02", "Performance", "Pattern generation end-to-end (AI + render)", "< 5s"],
        ["NFR-03", "Performance", "Search query response time", "< 200ms"],
        ["NFR-04", "Scalability", "auth-service horizontal scale", "10k req/min"],
        ["NFR-05", "Scalability", "pattern-service DynamoDB write throughput", "1000 WCU"],
        ["NFR-06", "Reliability", "System uptime SLA", "99.9%"],
        ["NFR-07", "Security", "All passwords hashed with bcrypt (cost 12)", "Mandatory"],
        ["NFR-08", "Security", "Access tokens expire in 15 minutes", "Mandatory"],
        ["NFR-09", "Security", "Refresh token rotation on every use", "Mandatory"],
        ["NFR-10", "Security", "All services behind API Gateway (no direct exposure)", "Mandatory"],
        ["NFR-11", "Accessibility", "WCAG 2.1 AA compliance for all MFE UIs", "Should Have"],
        ["NFR-12", "Observability", "Structured JSON logs (pino) on all services", "Must Have"],
        ["NFR-13", "Observability", "Health check endpoint on every service", "Must Have"],
    ],
    [0.8, 1.2, 3.2, 1.3]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("7. System Architecture")

heading2("7.1 Architecture Pattern")
body("PatternForge uses a Micro Frontend (MFE) + Microservices architecture hosted on AWS, with a Turborepo monorepo for unified tooling and CI/CD.")

heading2("7.2 Micro Frontend Apps")
add_table(
    ["App", "Port", "Framework", "Responsibility"],
    [
        ["shell", "3000", "Next.js 14", "Host container — composes all MFEs, global nav"],
        ["mfe-auth", "3001", "Next.js 14", "Login, register, profile, session management"],
        ["mfe-editor", "3002", "Next.js 14 + R3F", "AI prompt input + 3D pattern editor"],
        ["mfe-gallery", "3003", "Next.js 14", "Public pattern discovery, search, likes"],
        ["mfe-dashboard", "3004", "Next.js 14", "User stats, history, profile settings"],
    ],
    [1.2, 0.8, 1.5, 3.0]
)

heading2("7.3 Backend Microservices")
add_table(
    ["Service", "Port", "Database", "Responsibility"],
    [
        ["api-gateway", "4000", "—", "Routing, rate limiting, JWT validation at edge"],
        ["auth-service", "4001", "PostgreSQL", "User registration, login, token lifecycle"],
        ["pattern-service", "4002", "DynamoDB", "Pattern CRUD, versioning, high-throughput writes"],
        ["search-service", "4003", "OpenSearch", "Full-text + vector search across patterns"],
        ["ai-service", "4004", "—", "Claude API integration, prompt-to-parameters"],
        ["notification-service", "4005", "SQS / SES", "Async email notifications and event processing"],
    ],
    [1.5, 0.8, 1.3, 2.9]
)

heading2("7.4 Database Architecture Decision")
add_table(
    ["Database", "Owner Service", "Why This Database"],
    [
        ["PostgreSQL", "auth-service", "Relational data (users, tokens) requiring strong consistency and ACID transactions"],
        ["DynamoDB", "pattern-service", "High-throughput pattern writes/reads at scale; single-digit millisecond latency; auto-scaling"],
        ["OpenSearch", "search-service", "Full-text indexing, relevance ranking, and vector similarity for pattern discovery"],
    ],
    [1.3, 1.5, 3.7]
)
body("Each database is owned exclusively by one service. No cross-service database access is permitted.")

heading2("7.5 Shared Packages")
add_table(
    ["Package", "Contents"],
    [
        ["@patternforge/shared-types", "TypeScript types — User, Pattern, PatternParameters, ApiResponse"],
        ["@patternforge/shared-ui", "React component library — Button, Input, Card (Tailwind-based)"],
        ["@patternforge/shared-utils", "Utility functions — cn(), formatDate()"],
        ["@patternforge/shared-config", "Base configs — TypeScript, ESLint, Prettier"],
    ],
    [2.5, 4.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("8. Technology Stack")

add_table(
    ["Layer", "Technology", "Justification"],
    [
        ["Frontend Framework", "Next.js 14 (App Router)", "SSR, SSG, file-based routing, industry standard"],
        ["UI Styling", "Tailwind CSS", "Utility-first, consistent dark theme, rapid iteration"],
        ["3D Rendering", "React Three Fiber + Three.js", "Declarative Three.js in React, large ecosystem"],
        ["State Management", "Zustand", "Lightweight, boilerplate-free, editor state"],
        ["Animation", "Framer Motion", "Production-quality animations in gallery MFE"],
        ["Charts", "Recharts", "Accessible, composable charting for dashboard"],
        ["Forms", "React Hook Form + Zod", "Type-safe validation, minimal re-renders"],
        ["Backend Framework", "Fastify", "Faster than Express, TypeScript-first, JSON schema"],
        ["ORM", "Prisma", "Type-safe PostgreSQL access, auto-migrations"],
        ["AI", "Anthropic Claude API", "Best-in-class instruction following for structured output"],
        ["Monorepo", "Turborepo", "Incremental builds, task caching, workspace management"],
        ["Infrastructure", "AWS CDK (TypeScript)", "IaC, type-safe AWS resource definitions"],
        ["CI/CD", "GitHub Actions", "Native GitHub integration, free for public repos"],
        ["Container Runtime", "Docker + Colima", "Local development parity with production"],
    ],
    [1.8, 2.2, 2.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. DATA MODELS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("9. Data Models")

heading2("9.1 User (PostgreSQL)")
add_table(
    ["Field", "Type", "Description"],
    [
        ["id", "String (cuid)", "Primary key"],
        ["email", "String (unique)", "User's email address"],
        ["username", "String (unique)", "Public handle (alphanumeric + underscore)"],
        ["displayName", "String", "Full display name"],
        ["avatarUrl", "String?", "Optional profile image URL"],
        ["passwordHash", "String", "bcrypt hash (cost 12)"],
        ["createdAt", "DateTime", "Account creation timestamp"],
        ["updatedAt", "DateTime", "Last update timestamp"],
    ],
    [1.5, 1.5, 3.5]
)

heading2("9.2 Pattern (DynamoDB)")
add_table(
    ["Attribute", "Type", "Description"],
    [
        ["id (PK)", "String", "Pattern unique identifier"],
        ["userId (SK)", "String", "Owner user ID"],
        ["title", "String", "Pattern title"],
        ["description", "String", "Human-readable description"],
        ["prompt", "String", "Original AI prompt"],
        ["parameters", "Map", "Full PatternParameters JSON"],
        ["style", "String", "Pattern style (GSI partition key)"],
        ["isPublic", "Boolean", "Visibility flag"],
        ["tags", "List<String>", "Searchable tags"],
        ["likes", "Number", "Like count (atomic increment)"],
        ["thumbnailUrl", "String?", "S3 URL of preview image"],
        ["createdAt", "String (ISO)", "Creation timestamp"],
    ],
    [1.8, 1.3, 3.4]
)

heading2("9.3 PatternParameters Schema")
add_table(
    ["Field", "Type", "Range / Values"],
    [
        ["style", "Enum", "geometric | organic | fractal | noise | parametric"],
        ["complexity", "Number", "0.0 – 1.0"],
        ["scale", "Number", "0.1 – 10.0"],
        ["rotation", "Number", "0 – 360 degrees"],
        ["symmetry", "Number", "1 – 12 (fold symmetry)"],
        ["colorPalette.primary", "String", "Hex color"],
        ["colorPalette.secondary", "String", "Hex color"],
        ["colorPalette.accent", "String", "Hex color"],
        ["colorPalette.background", "String", "Hex color"],
        ["seed", "Number", "Integer — deterministic reproduction"],
        ["customParams", "Record", "Style-specific parameters"],
    ],
    [2.0, 1.2, 3.3]
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. API SPECIFICATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("10. API Specification")

heading2("10.1 Auth Service (port 4001)")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/auth/register", "None", "Register new user → returns tokens"],
        ["POST", "/auth/login", "None", "Login → returns tokens"],
        ["POST", "/auth/refresh", "None", "Rotate refresh token → new token pair"],
        ["POST", "/auth/logout", "None", "Invalidate refresh token"],
        ["GET", "/auth/me", "Bearer JWT", "Get current authenticated user"],
        ["GET", "/health", "None", "Service health check"],
    ],
    [0.8, 1.8, 1.0, 2.9]
)

heading2("10.2 Pattern Service (port 4002)")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/patterns", "Bearer JWT", "Create / save a new pattern"],
        ["GET", "/patterns/:id", "Optional", "Get pattern by ID"],
        ["PUT", "/patterns/:id", "Bearer JWT", "Update pattern (owner only)"],
        ["DELETE", "/patterns/:id", "Bearer JWT", "Delete pattern (owner only)"],
        ["GET", "/patterns/user/:userId", "Optional", "List patterns by user"],
        ["POST", "/patterns/:id/fork", "Bearer JWT", "Fork a public pattern"],
        ["POST", "/patterns/:id/like", "Bearer JWT", "Like / unlike a pattern"],
    ],
    [0.8, 1.8, 1.0, 2.9]
)

heading2("10.3 AI Service (port 4004)")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/ai/generate", "Bearer JWT", "Translate text prompt → PatternParameters"],
        ["POST", "/ai/refine", "Bearer JWT", "Refine existing parameters with a follow-up prompt"],
    ],
    [0.8, 1.8, 1.0, 2.9]
)

heading2("10.4 Search Service (port 4003)")
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["GET", "/search?q=&tags=&style=&page=", "None", "Full-text pattern search"],
        ["GET", "/search/similar/:id", "None", "Find visually similar patterns"],
        ["GET", "/search/trending", "None", "Top patterns by recent likes"],
    ],
    [0.8, 1.8, 1.0, 2.9]
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. IMPLEMENTATION ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("11. Implementation Roadmap")

add_table(
    ["Phase", "Milestone", "Key Deliverables", "Status"],
    [
        ["Phase 1", "Foundation", "Monorepo, MFE scaffold, microservice stubs, CI/CD", "✅ Complete"],
        ["Phase 1", "Auth Flow", "Register, login, JWT, refresh, logout, dashboard UI", "✅ Complete"],
        ["Phase 2", "AI + 3D Editor", "Claude API integration, React Three Fiber renderer, parameter editor", "🔄 Next"],
        ["Phase 2", "Pattern Service", "DynamoDB CRUD, save/fork/like, version history", "🔲 Planned"],
        ["Phase 3", "Search & Gallery", "OpenSearch indexing, full-text search, masonry gallery", "🔲 Planned"],
        ["Phase 3", "Dashboard", "Analytics UI, Recharts, profile edit", "🔲 Planned"],
        ["Phase 4", "Infrastructure", "AWS CDK stack, ECS deployment, RDS, production environment", "🔲 Planned"],
        ["Phase 4", "Polish & Launch", "WCAG audit, performance tuning, README architecture diagram", "🔲 Planned"],
    ],
    [0.8, 1.3, 3.4, 1.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# 12. RISKS & MITIGATIONS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("12. Risks & Mitigations")

add_table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Claude API latency > 5s", "Medium", "High", "Stream response; show skeleton render while AI processes"],
        ["DynamoDB cost overrun at scale", "Low", "Medium", "Use on-demand billing; set CloudWatch cost alarms"],
        ["OpenSearch index lag on saves", "Low", "Low", "Async indexing via SQS; search shows eventual consistency warning"],
        ["Three.js performance on low-end devices", "Medium", "Medium", "LOD (Level of Detail); cap complexity on mobile"],
        ["JWT access token leakage", "Low", "High", "Short 15-min TTL; refresh rotation; HTTPS only in production"],
        ["MFE bundle size too large", "Medium", "Medium", "Code splitting; lazy-load MFEs; share React via shell"],
    ],
    [2.2, 1.0, 0.8, 2.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 13. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("13. Glossary")

add_table(
    ["Term", "Definition"],
    [
        ["MFE", "Micro Frontend — independently deployable frontend application composed at runtime"],
        ["JWT", "JSON Web Token — signed token used for stateless authentication"],
        ["R3F", "React Three Fiber — React renderer for Three.js 3D graphics"],
        ["DynamoDB", "AWS managed NoSQL key-value and document database"],
        ["OpenSearch", "AWS managed search and analytics engine (Elasticsearch-compatible)"],
        ["Turborepo", "High-performance monorepo build system by Vercel"],
        ["AWS CDK", "AWS Cloud Development Kit — define cloud infrastructure in TypeScript"],
        ["SQS", "AWS Simple Queue Service — managed message queue for async processing"],
        ["SES", "AWS Simple Email Service — managed transactional email service"],
        ["PatternParameters", "Structured JSON schema defining all visual properties of a generated pattern"],
        ["NDCG", "Normalised Discounted Cumulative Gain — search relevance quality metric"],
        ["WCU", "Write Capacity Unit — DynamoDB throughput measurement"],
    ],
    [1.8, 4.7]
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/divyanksrivastava/patternforge/PatternForge_BRD.docx"
doc.save(output_path)
print(f"BRD saved → {output_path}")
